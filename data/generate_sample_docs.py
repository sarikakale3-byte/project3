"""
One-off generator for the synthetic InsuranceAssist knowledge base.

Run once:
    python -m data.generate_sample_docs

Output is committed to:
    data/knowledge_base/

This knowledge base supports:
- Claim submission
- Claim status tracking
- Policy coverage
- Claim rejection handling
- Reimbursement processing
- Insurance FAQs
"""

import csv
import os

from docx import Document
from fpdf import FPDF

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "knowledge_base")


def write_general_faq():
    html = """
<!DOCTYPE html>
<html>
<head>
<title>InsuranceAssist General FAQ</title>
</head>
<body>

<h1>InsuranceAssist - General Insurance FAQ</h1>

<h2>Claim Status Tracking</h2>
<p>
Customers can track their claims using the claim number and registered
mobile number through the customer portal, mobile application, or
customer support center.
</p>

<h2>Claim Settlement Timelines</h2>
<p>
Claims are generally processed within 7 to 15 working days after receipt
of complete documentation. Complex claims requiring investigation may
take additional time.
</p>

<h2>Customer Support</h2>
<p>
Customers can raise service requests via the mobile application, web
portal, or call the 24x7 support helpline. Complaints are acknowledged
within 24 hours and resolved within 7 working days.
</p>

</body>
</html>
"""

    path = os.path.join(OUT_DIR, "insurance_general_faq.html")

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)

    print("wrote", path)


def write_policy_coverage_faq():
    html = """
<!DOCTYPE html>
<html>
<head>
<title>Policy Coverage FAQ</title>
</head>
<body>

<h1>Policy Coverage and Exclusions</h1>

<h2>Health Insurance Coverage</h2>
<p>
Coverage typically includes hospitalization expenses, intensive care,
day-care procedures, ambulance charges, and pre/post hospitalization
expenses as defined by the policy document.
</p>

<h2>Motor Insurance Coverage</h2>
<p>
Coverage includes accidental damage, theft, fire, natural calamities,
and third-party liability as per policy terms.
</p>

<h2>Common Exclusions</h2>
<p>
Policies generally exclude cosmetic procedures, experimental treatments,
intentional self-harm, damages caused while driving under the influence,
and losses not specifically covered under policy terms.
</p>

</body>
</html>
"""

    path = os.path.join(OUT_DIR, "policy_coverage_faq.html")

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)

    print("wrote", path)


def write_claim_submission_guide():
    doc = Document()

    doc.add_heading(
        "Insurance Claim Submission Guide",
        level=1
    )

    doc.add_heading(
        "Health Insurance Claims",
        level=2
    )

    doc.add_paragraph(
        "Customers must submit a completed claim form along with "
        "hospital bills, discharge summary, prescriptions, "
        "diagnostic reports, and identity proof."
    )

    doc.add_heading(
        "Motor Insurance Claims",
        level=2
    )

    doc.add_paragraph(
        "Customers should report the incident immediately and "
        "provide photographs, vehicle registration certificate, "
        "driving license, repair estimates, and FIR when applicable."
    )

    doc.add_heading(
        "Life Insurance Claims",
        level=2
    )

    doc.add_paragraph(
        "Required documents include the claim form, policy document, "
        "death certificate, nominee identity proof, and bank account details."
    )

    doc.add_heading(
        "Travel Insurance Claims",
        level=2
    )

    doc.add_paragraph(
        "Claims may require travel tickets, passport copies, "
        "expense invoices, and supporting incident documentation."
    )

    path = os.path.join(
        OUT_DIR,
        "claim_submission_guide.docx"
    )

    doc.save(path)

    print("wrote", path)


def write_claim_rejection_policy():
    pdf = FPDF()
    
    # Set margins to ensure enough space for text rendering
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.set_top_margin(15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Reset x position to left margin
    pdf.set_x(15)

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(
        0,
        10,
        "Insurance Claim Rejection and Appeals Policy"
    )

    pdf.ln(5)

    sections = [
        (
            "Common Rejection Reasons",
            "Claims may be rejected due to policy exclusions, "
            "non-disclosure of material facts, policy lapse, "
            "fraudulent information, incomplete documentation, "
            "or incidents occurring during waiting periods."
        ),
        (
            "Document Verification",
            "Claims cannot proceed until all required documents "
            "have been submitted and verified."
        ),
        (
            "Appeal Process",
            "Customers may appeal rejected claims by providing "
            "additional documentation or clarification within "
            "30 calendar days of receiving the rejection notice."
        ),
        (
            "Investigation",
            "Certain claims may require detailed investigation "
            "before a final decision is reached."
        ),
        (
            "Grievance Escalation",
            "Unresolved concerns may be escalated to the "
            "Grievance Redressal Team through official support channels."
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_x(15)
        pdf.multi_cell(0, 8, title)

        pdf.set_font("Helvetica", "", 11)
        pdf.set_x(15)
        pdf.multi_cell(0, 6, body)

        pdf.ln(2)

    path = os.path.join(
        OUT_DIR,
        "claim_rejection_policy.pdf"
    )

    pdf.output(path)

    print("wrote", path)


def write_reimbursement_policy():
    pdf = FPDF()
    
    # Set margins to ensure enough space for text rendering
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)
    pdf.set_top_margin(15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Reset x position to left margin
    pdf.set_x(15)

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(
        0,
        10,
        "Claim Reimbursement and Settlement Policy"
    )

    pdf.ln(4)

    sections = [
        (
            "Reimbursement Claims",
            "Customers who receive treatment at a non-network "
            "facility may submit reimbursement claims after "
            "paying expenses directly."
        ),
        (
            "Settlement Timelines",
            "Most claims are settled within 7-15 working days "
            "after all required documents have been verified."
        ),
        (
            "Payment Methods",
            "Approved claim amounts are transferred directly "
            "to the customer's registered bank account."
        ),
        (
            "Status Updates",
            "Customers receive SMS, email, and portal updates "
            "throughout the claim lifecycle."
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_x(15)
        pdf.multi_cell(0, 8, title)

        pdf.set_font("Helvetica", "", 11)
        pdf.set_x(15)
        pdf.multi_cell(0, 6, body)

        pdf.ln(2)

    path = os.path.join(
        OUT_DIR,
        "reimbursement_policy.pdf"
    )

    pdf.output(path)

    print("wrote", path)


def write_claim_document_checklist():
    rows = [
        [
            "claim_type",
            "required_documents",
            "processing_time"
        ],
        [
            "Health Insurance",
            "Claim Form; Hospital Bills; Discharge Summary; Prescriptions",
            "7-15 Working Days"
        ],
        [
            "Motor Insurance",
            "Claim Form; RC; Driving License; Repair Estimate",
            "5-15 Working Days"
        ],
        [
            "Life Insurance",
            "Claim Form; Death Certificate; Policy Copy",
            "10-20 Working Days"
        ],
        [
            "Travel Insurance",
            "Tickets; Passport Copy; Expense Bills",
            "7-15 Working Days"
        ],
        [
            "Property Insurance",
            "Claim Form; Damage Assessment Report; Photos",
            "10-30 Working Days"
        ]
    ]

    path = os.path.join(
        OUT_DIR,
        "claim_documents_checklist.csv"
    )

    with open(
        path,
        "w",
        newline="",
        encoding="utf-8"
    ) as f:
        writer = csv.writer(f)
        writer.writerows(rows)

    print("wrote", path)


def main():
    os.makedirs(
        OUT_DIR,
        exist_ok=True
    )

    write_general_faq()
    write_policy_coverage_faq()
    write_claim_submission_guide()
    write_claim_rejection_policy()
    write_reimbursement_policy()
    write_claim_document_checklist()


if __name__ == "__main__":
    main()